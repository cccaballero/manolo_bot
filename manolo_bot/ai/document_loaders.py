import io
import logging
from collections.abc import Iterator
from typing import BinaryIO

from langchain_community.document_loaders.base import BaseBlobParser
from langchain_community.document_loaders.parsers import PyPDFParser
from langchain_community.document_loaders.parsers.txt import TextParser
from langchain_core.document_loaders import Blob
from langchain_core.documents import Document


class UnsupportedFileError(ValueError):
    """Exception raised when a file format is not supported."""

    pass


class DocxParser(BaseBlobParser):
    """
    Parser for DOCX files using python-docx.
    Follows the LangChain BaseBlobParser interface.
    We use this instead of MsWordParser to avoid the heavy 'unstructured' dependency.
    """

    def lazy_parse(self, blob: Blob) -> Iterator[Document]:
        import docx

        with blob.as_bytes_io() as file_like:
            doc = docx.Document(file_like)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            yield Document(page_content=text, metadata={"source": blob.source})


class DocumentLoader:
    """
    Utility class for extracting text from different document formats using LangChain native parsers.
    """

    _extensions = {"pdf": ["pdf"], "docx": ["docx"], "txt": ["txt", "md", "csv"]}
    SUPPORTED_EXTENSIONS = [ext for ext_list in _extensions.values() for ext in ext_list]

    def __init__(self):
        # mode="single" returns the whole document as one Document object
        self.pdf_parser = PyPDFParser(mode="single", pages_delimiter="\n")
        self.docx_parser = DocxParser()
        self.txt_parser = TextParser()

    @classmethod
    def validate_filename(cls, filename: str) -> None:
        """
        Validates if a filename has a supported extension.

        :param filename: The filename to validate.
        :raises UnsupportedFileError: If the extension is not supported.
        """
        if not filename or "." not in filename:
            raise UnsupportedFileError("File has no extension")

        extension = filename.split(".")[-1].lower()
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileError(f"Unsupported file extension: {extension}")

    def extract_text_from_pdf(self, file: BinaryIO) -> str:
        """
        Extracts text from a PDF file using PyPDFParser.
        """
        try:
            blob = Blob.from_data(file.read(), mime_type="application/pdf")
            docs = list(self.pdf_parser.lazy_parse(blob))
            return "\n".join([doc.page_content for doc in docs]).strip()
        except Exception as e:
            logging.error(f"Error extracting text from PDF: {e}")
            raise

    def extract_text_from_docx(self, file: BinaryIO) -> str:
        """
        Extracts text from a DOCX file using DocxParser.
        """
        try:
            blob = Blob.from_data(
                file.read(), mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            docs = list(self.docx_parser.lazy_parse(blob))
            return "\n".join([doc.page_content for doc in docs]).strip()
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {e}")
            raise

    def extract_text_from_txt(self, file: BinaryIO) -> str:
        """
        Extracts text from a TXT/MD/CSV file using TextParser.
        """
        try:
            blob = Blob.from_data(file.read(), mime_type="text/plain")
            docs = list(self.txt_parser.lazy_parse(blob))
            return "\n".join([doc.page_content for doc in docs]).strip()
        except Exception as e:
            logging.error(f"Error extracting text from TXT: {e}")
            raise

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Dispatcher method to extract text based on file extension.
        """
        # Ensure validation is performed (it will raise UnsupportedFileError if invalid)
        self.validate_filename(filename)

        file_like = io.BytesIO(file_content)
        extension = filename.split(".")[-1].lower()

        if extension in self._extensions["pdf"]:
            return self.extract_text_from_pdf(file_like)
        elif extension in self._extensions["docx"]:
            return self.extract_text_from_docx(file_like)
        else:
            # Must be txt, md, or csv based on validate_filename check
            return self.extract_text_from_txt(file_like)


def clean_text(text: str) -> str:
    """
    Basic text cleaning to reduce token usage.
    Removes multiple whitespaces and newlines.
    """
    import re

    # Replace multiple spaces with a single space
    text = re.sub(r" +", " ", text)
    # Replace multiple newlines with a double newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
